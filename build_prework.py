# -*- coding: utf-8 -*-
"""クライアント向け事前準備資料（Word課題シート＋メール文面テキスト）を生成"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ===== アプリと統一したカラーパレット =====
BRAND   = RGBColor(0x2f, 0x8f, 0x7f)
BRAND_D = RGBColor(0x22, 0x6b, 0x5f)
BRAND_L_HEX = "E7F4F1"
PALE_HEX    = "F4F7F6"
WARN_L_HEX  = "FBF2E3"
ACCENT_L_HEX = "ECECFB"
INK     = RGBColor(0x23, 0x32, 0x3a)
MUTED   = RGBColor(0x6b, 0x7c, 0x85)
WARN_D  = RGBColor(0xc9, 0x8a, 0x2b)

FONT_J = "游ゴシック"

def set_japanese_font(run, name=FONT_J):
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)

def add_run(p, text, *, size=11, bold=False, color=INK, font=FONT_J):
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color
    set_japanese_font(run, font)
    return run

def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)

def set_cell_borders(cell, color="DBE6E3", sz="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top","left","bottom","right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),"single")
        b.set(qn("w:sz"),sz)
        b.set(qn("w:color"),color)
        tcBorders.append(b)
    tc_pr.append(tcBorders)

def add_horizontal_band(doc, color_hex, height_pt=6):
    """細い色帯を1行のテーブルで描く"""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, color_hex)
    # 高さを小さく
    tr = cell._tc.getparent()
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(height_pt*20)))
    trHeight.set(qn("w:hRule"), "exact")
    trPr.append(trHeight)
    # 段落を空に
    cell.paragraphs[0].text = ""

# ============================================================
def build_word():
    doc = Document()

    # 余白
    for s in doc.sections:
        s.left_margin = Cm(2.2)
        s.right_margin = Cm(2.2)
        s.top_margin = Cm(2.0)
        s.bottom_margin = Cm(2.0)

    # 標準スタイル
    style = doc.styles["Normal"]
    style.font.name = FONT_J
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_J)

    # ── ブランド帯 ──
    add_horizontal_band(doc, "2f8f7f", height_pt=4)

    # ── タイトル ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p, "NLPコーチング セッション", size=10, color=MUTED)
    p = doc.add_paragraph()
    add_run(p, "事前ワークシート", size=24, bold=True, color=BRAND_D)

    p = doc.add_paragraph()
    add_run(p, "〜 今回のセッションで扱いたい課題を整理しましょう 〜", size=11, color=MUTED)

    doc.add_paragraph()

    # ── 趣旨 ──
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, BRAND_L_HEX)
    set_cell_borders(cell, color="CFE7E1")
    cell.text = ""
    p = cell.paragraphs[0]
    add_run(p, "📌 このシートについて", size=11, bold=True, color=BRAND_D)
    p2 = cell.add_paragraph()
    add_run(p2, (
        "セッションの限られた時間を、あなたの本当に大切な課題に使うためのワークシートです。"
        "今、頭の中にある『気になっていること』『手放したいこと』『叶えたいこと』を、"
        "思いつくままに書き出してみてください。\n"
        "正解はありません。完璧でなくて大丈夫。書き出すことそのものが、最初の整理になります。"
    ), size=10.5)

    doc.add_paragraph()

    # ── 記入欄ヘッダー ──
    p = doc.add_paragraph()
    add_run(p, "✏️ 解決したい課題リスト（5〜10個）", size=14, bold=True, color=BRAND_D)
    p = doc.add_paragraph()
    add_run(p, "下の欄に、セッションで扱いたい課題を箇条書きで書き出してください。", size=10.5, color=MUTED)

    # 例
    example_table = doc.add_table(rows=1, cols=1)
    ecell = example_table.rows[0].cells[0]
    set_cell_bg(ecell, ACCENT_L_HEX)
    set_cell_borders(ecell, color="DCDCF7")
    ecell.text = ""
    ep = ecell.paragraphs[0]
    add_run(ep, "💡 書き方の例", size=10.5, bold=True, color=RGBColor(0x6b,0x6a,0xd6))
    examples = [
        "人前で話すと緊張してしまい、本来の力が出せない",
        "やるべきことをつい後回しにしてしまう",
        "上司との関係がぎくしゃくしていて、毎日ストレスを感じる",
        "自分に自信が持てず、新しいことに踏み出せない",
        "将来やりたいことが見えず、もやもやしている",
        "パートナーとの関係を、もっと温かく深めたい",
        "怒りや不安などの感情に振り回されることがある",
    ]
    for ex in examples:
        ep2 = ecell.add_paragraph()
        add_run(ep2, f"・{ex}", size=10.5, color=INK)
    ep3 = ecell.add_paragraph()
    add_run(ep3, "（仕事、家族、お金、健康、人間関係、自分自身…どんな領域でもOK）",
            size=10, color=MUTED)

    doc.add_paragraph()

    # ── 入力欄テーブル（10行） ──
    table = doc.add_table(rows=11, cols=2)
    table.autofit = False
    # 列幅
    for row in table.rows:
        row.cells[0].width = Cm(1.5)
        row.cells[1].width = Cm(14.5)

    # ヘッダー
    hdr = table.rows[0]
    set_cell_bg(hdr.cells[0], "2f8f7f")
    set_cell_bg(hdr.cells[1], "2f8f7f")
    set_cell_borders(hdr.cells[0], color="226B5F")
    set_cell_borders(hdr.cells[1], color="226B5F")
    hdr.cells[0].text = ""
    hdr.cells[1].text = ""
    p = hdr.cells[0].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "#", size=11, bold=True, color=RGBColor(0xff,0xff,0xff))
    p = hdr.cells[1].paragraphs[0]
    add_run(p, "セッションで扱いたい課題", size=11, bold=True, color=RGBColor(0xff,0xff,0xff))

    for i in range(1, 11):
        row = table.rows[i]
        # 行の高さ
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement("w:trHeight")
        trHeight.set(qn("w:val"), "560")
        trPr.append(trHeight)

        n_cell = row.cells[0]
        b_cell = row.cells[1]
        set_cell_bg(n_cell, PALE_HEX)
        set_cell_borders(n_cell)
        set_cell_borders(b_cell)
        n_cell.text = ""
        b_cell.text = ""
        p = n_cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, str(i), size=12, bold=True, color=BRAND_D)
        # 入力欄は空に（罫線だけ）

    doc.add_paragraph()

    # ── 望む状態セクション ──
    p = doc.add_paragraph()
    add_run(p, "🎯 今回扱いたい『1つ』をあらかじめ選べる方は", size=12, bold=True, color=BRAND_D)
    p = doc.add_paragraph()
    add_run(p, "上のリストの中から、特に今回のセッションで扱いたい課題が決まっていれば、下にお書きください（任意）。", size=10.5, color=MUTED)

    fields = [
        ("今回扱いたい課題（リストから1つ）", 1, BRAND_L_HEX),
        ("現状（Present State）今、どんな状態ですか？", 3, PALE_HEX),
        ("望む状態（Desired State）どうなっていたいですか？", 3, PALE_HEX),
        ("困りごとの強さ（10段階）／ どのくらい辛い・気になる？", 1, BRAND_L_HEX),
        ("セッション終了後にどうなっていたいか？（今日のゴール）", 2, BRAND_L_HEX),
    ]
    for label, lines, bg in fields:
        tbl = doc.add_table(rows=1+lines, cols=1)
        # ラベル
        lcell = tbl.rows[0].cells[0]
        set_cell_bg(lcell, "2f8f7f")
        set_cell_borders(lcell, color="226B5F")
        lcell.text = ""
        lp = lcell.paragraphs[0]
        add_run(lp, label, size=10.5, bold=True, color=RGBColor(0xff,0xff,0xff))
        # 入力欄
        for j in range(lines):
            ic = tbl.rows[1+j].cells[0]
            set_cell_bg(ic, "ffffff")
            set_cell_borders(ic)
            ic.text = ""
            # 行高
            tr = tbl.rows[1+j]._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement("w:trHeight")
            trHeight.set(qn("w:val"), "560")
            trPr.append(trHeight)
        doc.add_paragraph()

    # ── 注意書き ──
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, WARN_L_HEX)
    set_cell_borders(cell, color="EED9B0")
    cell.text = ""
    p = cell.paragraphs[0]
    add_run(p, "🌿 心構え", size=11, bold=True, color=WARN_D)
    p2 = cell.add_paragraph()
    add_run(p2, (
        "・無理に多く書こうとしなくて大丈夫です。3つでも、1つでも構いません。\n"
        "・「ささいなことかも」と思っても、気になるならぜひ書いてみてください。\n"
        "・書きながら涙が出たり、感情が動くことがあります。それも大切なサインです。\n"
        "・このシートはセッション中、私（コーチ）が一緒に拝見し、扱う課題を一緒に選んでいきます。"
    ), size=10.5, color=INK)

    doc.add_paragraph()
    doc.add_paragraph()

    # ── フッター ──
    add_horizontal_band(doc, "DBE6E3", height_pt=2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "NLP Coaching ／ Pre-Session Worksheet", size=9, color=MUTED)

    doc.save("nlp_prework_sheet.docx")
    print("nlp_prework_sheet.docx を生成しました")


def build_email_text():
    """セッション前に送るメール文面（テキスト）"""
    text = """件名：【NLPコーチングセッション】事前のご準備のお願い

○○ 様

このたびはNLPコーチングセッションをお申し込みいただき、
誠にありがとうございます。

セッションの時間をあなたにとって最大限に有意義なものにするため、
当日までに以下のご準備をお願いいたします。
（10〜15分ほどで完了します）


■ 事前にお考えいただきたいこと

セッションで扱いたい「課題」を、5〜10個ほど箇条書きで
書き出してきてください。

今、頭の中にある『気になっていること』『手放したいこと』
『叶えたいこと』を、思いつくままに書き出していただければ十分です。
仕事・家族・お金・健康・人間関係・自分自身…どんな領域でも構いません。


■ 課題の書き方の例

・人前で話すと緊張してしまい、本来の力が出せない
・やるべきことをつい後回しにしてしまう
・上司との関係がぎくしゃくしていて、毎日ストレスを感じる
・自分に自信が持てず、新しいことに踏み出せない
・将来やりたいことが見えず、もやもやしている
・パートナーとの関係を、もっと温かく深めたい
・怒りや不安などの感情に振り回されることがある


■ ご記入のヒント

・完璧に書こうとしなくて大丈夫です。
　書き出すこと自体が、最初の整理になります。
・「ささいなことかも」と思っても、気になるならぜひ書いてみてください。
・どれか1つだけでも構いません。当日一緒に深めていきましょう。


■ ご準備の方法

下記いずれかの方法でお持ちいただければ大丈夫です。

  ① 添付のWordファイル（事前ワークシート）に記入してお持ちいただく
  ② このメールに直接ご返信いただく（または当日コピペでお持ちください）
  ③ 紙やメモ帳にご自身で書いてお持ちいただく


■ セッション当日について

日時：　　年　　月　　日（　）　　：　　〜
場所／方法：

当日は、お持ちいただいた課題リストの中から
今回特に扱いたい1つを一緒に選び、深めていきます。


何かご不明な点がございましたら、お気軽にご連絡ください。
当日お会いできることを楽しみにしています。


────────────────────────────
NLPコーチ
（あなたのお名前／屋号）
────────────────────────────
"""
    with open("nlp_prework_email.txt", "w", encoding="utf-8") as f:
        f.write(text)
    print("nlp_prework_email.txt を生成しました")


if __name__ == "__main__":
    build_word()
    build_email_text()
